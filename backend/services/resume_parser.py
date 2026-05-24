import os
from pypdf import PdfReader
import docx

class ResumeParserService:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extracts text from a PDF file."""
        try:
            reader = PdfReader(file_path)
            text_parts = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            full_text = "\n".join(text_parts)
            cleaned_text = ResumeParserService.clean_text(full_text)
            if not cleaned_text.strip():
                raise ValueError("The PDF file does not contain any readable text.")
            return cleaned_text
        except Exception as e:
            raise ValueError(f"Failed to parse PDF resume: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extracts text from a DOCX file."""
        try:
            doc = docx.Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)

            full_text = "\n".join(text_parts)
            cleaned_text = ResumeParserService.clean_text(full_text)
            if not cleaned_text.strip():
                raise ValueError("The DOCX file does not contain any readable text.")
            return cleaned_text
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX resume: {str(e)}")

    @staticmethod
    def clean_text(text: str) -> str:
        """Cleans up formatting and normalizes whitespaces in extracted text."""
        # Replace multiple consecutive spaces with a single space
        lines = [line.strip() for line in text.splitlines()]
        # Remove empty lines
        cleaned_lines = [line for line in lines if line]
        return "\n".join(cleaned_lines)

    @staticmethod
    def parse_resume(file_path: str) -> str:
        """Orchestrator to detect extension and parse appropriately."""
        if not os.path.exists(file_path):
            raise FileNotFoundError("Resume file not found at the specified path.")

        _, ext = os.path.splitext(file_path.lower())
        if ext == ".pdf":
            return ResumeParserService.extract_text_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return ResumeParserService.extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
